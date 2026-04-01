from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Last-Mile Delivery Failure Diagnostics System', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('Insight Report — January 2026', 2)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Meta info
doc.add_paragraph('Dataset: NYC Taxi Trip Records 2026')
doc.add_paragraph('Records Analyzed: 3,504,776 deliveries')
doc.add_paragraph('Tools: Python · MySQL · Power BI · SQL')
doc.add_paragraph('Prepared by: Akshay Dusane')

doc.add_paragraph('')

# Executive Summary
doc.add_heading('Executive Summary', 1)
doc.add_paragraph(
    'Analysis of 3.5 million last-mile deliveries revealed an overall '
    'SLA breach rate of 18.34% — meaning 642,723 deliveries failed to meet '
    'the promised delivery window. The network faces a systemic capacity '
    'problem concentrated in specific zones, routes, and time windows.'
)

# Key Findings
doc.add_heading('Key Findings', 1)

doc.add_heading('Finding 1 — Network-Wide SLA Performance', 2)
findings1 = [
    'Total deliveries analyzed: 3,504,776',
    'On-time delivery rate: 81.66%',
    'Total SLA breaches: 642,723',
    'Average delivery time: 17.2 minutes',
    'Average delivery cost: $21.03'
]
for f in findings1:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('Finding 2 — Critical Route Failures', 2)
findings2 = [
    'Route 76_140 has a 100% SLA breach rate across 52 deliveries',
    'Route 76_181 has a 96.88% breach rate across 64 deliveries',
    'Top 20 routes all fall in the critical quartile',
    'This indicates a systemic network problem, not isolated failures'
]
for f in findings2:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('Finding 3 — Peak Hour Impact', 2)
findings3 = [
    'Peak weekday is the worst segment at 21.41% breach rate',
    'Off-peak weekend is the healthiest at 11.06% breach rate',
    'Peak hours have nearly 2x the breach rate of off-peak hours',
    '1,078,116 deliveries occur during peak weekday slots'
]
for f in findings3:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('Finding 4 — Zone-Level Capacity Crisis', 2)
findings4 = [
    'Zone 117: 72.47% breach rate across 1,108 dispatches',
    'Zone 86: 70.29% breach rate across 1,067 dispatches',
    'Zone 222: 66.63% breach rate across 1,046 dispatches',
    '7 origin zones identified as critically overloaded'
]
for f in findings4:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('Finding 5 — Hourly Anomalies', 2)
findings5 = [
    'Thursday 3PM has the highest breach rate at 32.44%',
    'Friday 3PM follows at 32.46%',
    'Thursday midnight shows anomalous 14.67% breach rate',
    'Afternoon slots (2PM-5PM) consistently show 25-32% breach rates'
]
for f in findings5:
    doc.add_paragraph(f, style='List Bullet')

# Recommendations
doc.add_heading('Recommendations', 1)

doc.add_heading('Recommendation 1 — Zone Capacity Reallocation', 2)
doc.add_paragraph(
    'Immediately add dispatch capacity to Zones 117, 86, and 222. '
    'These three zones alone account for disproportionate breach volume. '
    'Reallocating resources from underutilized zones could reduce '
    'network-wide breach rate by an estimated 15-20%.'
)
doc.add_paragraph('Owner: Operations Leadership')
doc.add_paragraph('Timeline: Immediate — within 2 weeks')

doc.add_heading('Recommendation 2 — Peak Hour Surge Staffing', 2)
doc.add_paragraph(
    'Introduce surge staffing protocol for peak weekday slots '
    '(7-10AM and 5-8PM Monday-Friday). Current breach rate of 21.41% '
    'during these windows is nearly double the off-peak weekend rate of 11.06%.'
)
doc.add_paragraph('Owner: Workforce Planning Team')
doc.add_paragraph('Timeline: Next scheduling cycle')

doc.add_heading('Recommendation 3 — Thursday/Friday Afternoon Audit', 2)
doc.add_paragraph(
    'Conduct an operational audit of Thursday and Friday afternoon slots '
    '(2PM-5PM) where breach rates consistently exceed 30%. Investigate '
    'whether this reflects driver shortage, route complexity, or traffic '
    'congestion patterns.'
)
doc.add_paragraph('Owner: Route Optimization Team')
doc.add_paragraph('Timeline: Within 30 days')

# Conclusion
doc.add_heading('Conclusion', 1)
doc.add_paragraph(
    'The data reveals that last-mile delivery failures are not random — '
    'they are concentrated in predictable zones, routes, and time windows. '
    'Targeted interventions in the top 7 overloaded zones and peak weekday '
    'staffing could recover an estimated 15-20% improvement in on-time '
    'delivery rate, directly impacting customer satisfaction and operational costs.'
)

# Footer
doc.add_paragraph('')
doc.add_paragraph(
    'Analysis conducted using MySQL (8 advanced SQL queries) and Power BI '
    'dashboard on NYC TLC Trip Record Data, January 2026.'
)

# Save
doc.save('docs/insight_report.docx')
print("Report saved to docs/insight_report.docx")